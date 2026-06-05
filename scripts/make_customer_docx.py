from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

VERSION = date.today().strftime("%Y-%m-%d")
OUTPUT  = f"/Users/sbilgram/Desktop/SF_Notes_Attachments_Assessment_{VERSION}.docx"
APP_URL = "https://sf-notes-migration.onrender.com"

doc = Document()
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

NAVY     = RGBColor(0x1A, 0x2F, 0x4E)
BLUE     = RGBColor(0x24, 0x78, 0xDB)
DARK     = RGBColor(0x2C, 0x3E, 0x50)
GREY     = RGBColor(0x7F, 0x8C, 0x8D)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
GREEN    = RGBColor(0x27, 0xAE, 0x60)
RED      = RGBColor(0xC0, 0x39, 0x2B)
ORANGE   = RGBColor(0xE6, 0x7E, 0x22)
TH_BG    = RGBColor(0x1A, 0x2F, 0x4E)
ROW_ALT  = RGBColor(0xF8, 0xF9, 0xFA)


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)


def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)


def para(text, color=None, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)


def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)


def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = GREY
    run.font.size = Pt(9)


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(9)
    run.font.color.rgb = GREEN
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F4F4F4')
    pPr.append(shd)


def rule():
    p = doc.add_paragraph('─' * 72)
    p.runs[0].font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    p.runs[0].font.size = Pt(8)


def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], TH_BG)
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
    for r_idx, row_data in enumerate(rows):
        cells = t.rows[r_idx + 1].cells
        for c_idx, text in enumerate(row_data):
            if r_idx % 2 == 1:
                set_cell_bg(cells[c_idx], ROW_ALT)
            run = cells[c_idx].paragraphs[0].add_run(text)
            run.font.size = Pt(9)
    doc.add_paragraph()


# ── Document ──────────────────────────────────────────────────────────────────

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
r = p.add_run('SF Notes & Attachments')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
r2 = p2.add_run('Phase 1 — Inventory & Assessment')
r2.font.size = Pt(14)
r2.font.color.rgb = DARK

p3 = doc.add_paragraph()
r3 = p3.add_run(f'Version {VERSION}')
r3.font.size = Pt(10)
r3.font.color.rgb = GREY
r3.italic = True

rule()

# Overview
h1('Overview')
para(
    'SF Notes & Attachments Phase 1 Inventory & Assessment is a read-only diagnostic tool '
    'that connects securely to your Salesforce org and produces a complete inventory of legacy '
    'Notes and Attachments. The assessment identifies volume, size, flags, and breakdown by '
    'parent object — giving you the information needed to plan a migration to modern Salesforce Files.'
)
para('All checks are read-only. No data is created, modified, or deleted in your org.')

rule()

# What the app checks
h1('What the App Checks')
para('The assessment runs five sections in parallel and typically completes in under two minutes.')

h2('1 — Legacy Notes')
bullet('Total count of legacy Notes in the org')
bullet('Count of Private Notes', bold_prefix='Private Notes: ')
note('Private Notes have no direct equivalent in modern Files. These require a handling decision before any migration — they can be skipped or migrated as internal-only records.')
bullet('Count of Notes with body exceeding 32KB', bold_prefix='Large Notes (>32KB): ')
note('ContentNote (the modern equivalent) has a 32KB body limit. Notes exceeding this size will be flagged at migration time.')
bullet('Breakdown of Notes by parent object type (Account, Contact, Case, Opportunity, etc.)')

h2('2 — Legacy Attachments')
bullet('Total count of legacy Attachments in the org')
bullet('Total combined size of all Attachments')
bullet('Count of Attachments exceeding 25MB', bold_prefix='Large Attachments (>25MB): ')
note('The Salesforce REST API has a 25MB per-file limit. Attachments exceeding this size cannot be migrated via the standard API and will require a separate approach.')
bullet('Breakdown of Attachments by parent object type')

h2('3 — Modern Salesforce Files (Baseline)')
bullet('Total count of existing ContentDocuments (modern Files already in the org)')
bullet('Total storage currently consumed by modern Files')
note('This baseline shows how much of the migration has already been completed and how much storage will be consumed after a full migration.')

h2('4 — Org Limits')
bullet('File Storage — used vs. total available, flagged at 65% (warning) and 85% (critical)')
bullet('Data Storage — used vs. total available')
bullet('API Requests remaining for the day vs. daily limit')
note('Storage headroom is checked before recommending a migration. If file storage is near capacity, the org may need a storage upgrade before migrating all attachments.')

h2('5 — Creation Controls')
bullet('Whether the org-level setting still allows new legacy Notes to be created')
bullet('Whether the org-level setting still allows new legacy Attachments to be created')
bullet('Every Profile that still grants Create permission on Notes or Attachments')
bullet('Every Permission Set that still grants Create permission on Notes or Attachments')
note('Identifying open creation permissions is essential to prevent the legacy record count from growing during or after a migration project.')

rule()

# Exports
h1('Exports')
para('After the assessment completes, click Preview & Export to open a full-screen formatted report. From there:')
bullet('Save / Print PDF — opens your browser\'s print dialog. Select "Save as PDF" to download.')
bullet('Export Excel — downloads a .xlsx workbook with one tab per section (Summary, Notes, Attachments, Org Limits, Creation Controls).')

rule()

# Privacy
h1('Data & Privacy')
para('All assessment data is ephemeral — nothing is written to a database or stored server-side beyond your active session.')

table(
    ['Data', 'Where it lives', 'Lifetime'],
    [
        ['OAuth tokens', 'Server memory (express-session)', '1 hour'],
        ['Assessment results', 'Browser tab (React state)', 'Until page refresh'],
        ['Credentials (if remembered)', 'Browser localStorage', 'Until you clear it'],
        ['Exported files (PDF / Excel)', 'Your local Downloads folder', 'Permanent'],
    ]
)
para('Nothing is written back to Salesforce. No external database is used.')

rule()

# Setup
h1('Setup: Connecting to Your Org')
para(
    'The app uses a Salesforce Connected App or External Client App for secure OAuth authentication. '
    'You will need to create one in your org — this is a one-time setup per org and takes approximately 5 minutes.'
)

h2('Step 1 — Determine your org type')
para('Go to Setup in your org and search for "External Client Apps" in the Quick Find box.')
bullet('If External Client Apps appears in the menu → use Option A below (newer orgs, Spring \'25+)')
bullet('If it does not appear → use Option B below (classic Connected App)')

h2('Step 2A — Create an External Client App (newer orgs)')
numbered('Log in as an Administrator → Setup → External Client Apps → New')
numbered('Fill in: Label = SF Notes Migration Assessment, Contact Email = your email')
numbered('Under OAuth Settings, check Enable OAuth')
numbered('Set Callback URL to:')
code(f'{APP_URL}/auth/callback')
numbered('Under OAuth Scopes, add:')
bullet('Manage user data via APIs (api)')
bullet('Perform requests at any time (refresh_token, offline_access)')
numbered('Uncheck "Require Proof Key for Code Exchange (PKCE)" — leave it disabled')
numbered('Click Save — wait 10 minutes for Salesforce to activate it')
numbered('Go back to the app → View Consumer Details → copy Consumer Key and Consumer Secret')

h2('Step 2B — Create a Connected App (older orgs)')
numbered('Log in as an Administrator → Setup → App Manager → New Connected App')
numbered('Fill in: Connected App Name = SF Notes Migration Assessment, Contact Email = your email')
numbered('Check Enable OAuth Settings')
numbered('Set Callback URL to:')
code(f'{APP_URL}/auth/callback')
numbered('Under Selected OAuth Scopes, add:')
bullet('Manage user data via APIs (api)')
bullet('Perform requests at any time (refresh_token, offline_access)')
numbered('Uncheck "Require Proof Key for Code Exchange (PKCE)"')
numbered('Click Save — wait 10 minutes for Salesforce to activate it')
numbered('Go back to the Connected App → Manage Consumer Details → copy Consumer Key and Consumer Secret')
numbered('Click Manage → Edit Policies → set Permitted Users to "All users may self-authorize" → Save')

h2('Step 3 — Find your Org URL')
para('Use your org\'s My Domain URL — not the Lightning UI URL.')

table(
    ['Org Type', 'Example URL format'],
    [
        ['Production', 'https://yourorg.my.salesforce.com'],
        ['Sandbox', 'https://yourorg--sandboxname.sandbox.my.salesforce.com'],
        ['Trailhead Playground', 'https://yourorg-dev-ed.trailblaze.my.salesforce.com'],
    ]
)
note('Do not use .lightning.force.com URLs — these will not work. Use the My Domain format above.')

h2('Step 4 — Run the Assessment')
numbered(f'Open {APP_URL}')
numbered('Enter your Org URL, Client ID (Consumer Key), and Client Secret (Consumer Secret)')
numbered('Click Connect to Salesforce and log in when prompted')
numbered('Click Run Assessment')
numbered('Once all five sections complete, click Preview & Export to view and download the report')

rule()

# Common errors
h1('Common Errors')

table(
    ['Error', 'Cause', 'Fix'],
    [
        ['redirect_uri_mismatch', 'Callback URL in Connected App does not match', f'Update Callback URL to exactly: {APP_URL}/auth/callback — save and wait 10 minutes'],
        ['missing required code challenge', 'PKCE is enabled on the Connected App', 'Go to Setup → uncheck "Require Proof Key for Code Exchange (PKCE)" → save and retry'],
        ['Authentication failed', 'Wrong Client ID or Client Secret', 'Go to Connected App → Manage Consumer Details → copy the exact Consumer Key and Secret'],
        ['Page will not load', 'Corporate firewall / Zscaler blocking onrender.com', 'Ask IT to whitelist sf-notes-migration.onrender.com, or run on a mobile hotspot'],
        ['App is slow to load (first visit)', 'Free tier hosting — app sleeps after inactivity', 'Wait ~30 seconds for the app to wake up, then reload'],
    ]
)

rule()

# Footer
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
r = p.add_run(f'SF Notes & Attachments Phase 1 Inventory & Assessment  |  Version {VERSION}')
r.font.size = Pt(8)
r.font.color.rgb = GREY
r.italic = True

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
